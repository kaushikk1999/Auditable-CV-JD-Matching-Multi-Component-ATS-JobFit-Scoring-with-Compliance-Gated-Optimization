
import sys
import io
from pathlib import Path
import pytest
from unittest.mock import MagicMock, patch

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from modules.parsers import CVParser, JDParser

def test_parse_txt():
    """Test TXT extraction."""
    content = b"John Doe\nSoftware Engineer"
    result = CVParser.parse(content, ".txt")
    assert "John Doe" in result
    assert "Software Engineer" in result

def test_parse_docx():
    """Test DOCX extraction using a generated file."""
    from docx import Document
    
    # Create a real docx in memory
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Data Scientist")
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    
    result = CVParser.parse(docx_bytes, ".docx")
    assert "Jane Doe" in result
    assert "Data Scientist" in result

@patch("modules.parsers.pdfplumber")
def test_parse_pdf(mock_pdfplumber):
    """Test PDF extraction using mocked pdfplumber."""
    # Setup mock
    mock_page1 = MagicMock()
    mock_page1.extract_text.return_value = "Alice Smith"
    
    mock_page2 = MagicMock()
    mock_page2.extract_text.return_value = "Project Manager"
    
    mock_pdf = MagicMock()
    mock_pdf.pages = [mock_page1, mock_page2]
    
    mock_pdf_context = MagicMock()
    mock_pdf_context.__enter__.return_value = mock_pdf
    mock_pdfplumber.open.return_value = mock_pdf_context
    
    content = b"fake pdf content"
    result = CVParser.parse(content, ".pdf")
    
    assert "Alice Smith" in result
    assert "Project Manager" in result

def test_clean_jd_text():
    """Test JD cleaning logic."""
    raw_jd = """
    Senior   Developer
    
    Required Skills:
    - Python
    
    """
    cleaned = JDParser.clean_text(raw_jd)
    
    # Should remove extra blank lines and leading/trailing whitespace
    assert "Senior   Developer" in cleaned  # Internal spaces might be preserved depending on logic
    assert "Required Skills:" in cleaned
    assert "- Python" in cleaned
    assert "\n\n" not in cleaned  # Assuming clean_text joins with single newline
    
    # Verify exact structure based on implementation:
    # lines = [line.strip() for line in jd_text.split('\n') if line.strip()]
    # return "\n".join(lines)
    expected = "Senior   Developer\nRequired Skills:\n- Python"
    assert cleaned == expected

def test_parse_unsupported_format():
    """Test error handling for unsupported formats."""
    with pytest.raises(ValueError, match="Unsupported file format"):
        CVParser.parse(b"content", ".xyz")
