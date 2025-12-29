
import os
from docx import Document
from reportlab.pdfgen import canvas

def create_txt():
    with open("dummy_cv.txt", "w") as f:
        f.write("John Doe\nSoftware Engineer\nSummary: Experienced Python developer.")
    print("Created dummy_cv.txt")

def create_docx():
    doc = Document()
    doc.add_heading('Jane Doe', 0)
    doc.add_paragraph('Data Scientist')
    doc.add_paragraph('Summary: Expert in Machine Learning and AI.')
    doc.save('dummy_cv.docx')
    print("Created dummy_cv.docx")

def create_pdf():
    c = canvas.Canvas("dummy_cv.pdf")
    c.drawString(100, 750, "Alice Smith")
    c.drawString(100, 730, "Project Manager")
    c.drawString(100, 710, "Summary: Skilled in Agile and Scrum.")
    c.showPage() # Page 1
    c.drawString(100, 750, "Page 2 Content")
    c.save()
    print("Created dummy_cv.pdf")

if __name__ == "__main__":
    create_txt()
    create_docx()
    try:
        create_pdf()
    except ImportError:
        print("ReportLab not installed, skipping PDF creation. Will try to find an existing PDF or skip PDF test.")
