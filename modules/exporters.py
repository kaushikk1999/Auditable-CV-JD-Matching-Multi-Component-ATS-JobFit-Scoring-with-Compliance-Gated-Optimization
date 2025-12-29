from typing import Dict
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import json
import markdown2
from datetime import datetime

class PDFExporter:
    """Export CV to ATS-friendly PDF."""
    
    def export(self, cv_text: str, output_path: Path):
        """
        Generate PDF with ATS-friendly formatting.
        
        - Single column
        - Black text only
        - Standard fonts (Arial/Helvetica)
        - No images, tables, or complex layouts
        """
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Use standard font
        pdf.set_font("Arial", size=11)
        
        # Split into lines
        lines = cv_text.split('\n')
        
        for line in lines:
            if not line.strip():
                pdf.ln(5)  # Blank line spacing
                continue
            
            # Sanitize for core fonts (Latin-1 only)
            safe_line = line.replace('–', '-').replace('•', '-').replace('“', '"').replace('”', '"').replace("’", "'")

            # Headers (all caps)
            if line.isupper() and len(line) < 50:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 6, txt=safe_line, ln=True)
                pdf.set_font("Arial", size=11)
            # Bullets
            elif line.strip().startswith('•'):
                pdf.multi_cell(0, 5, txt=safe_line, new_x="LMARGIN", new_y="NEXT")
            # Regular text
            else:
                pdf.multi_cell(0, 5, txt=safe_line, new_x="LMARGIN", new_y="NEXT")
        
        pdf.output(str(output_path))

class DOCXExporter:
    """Export CV to Microsoft Word format."""
    
    def export(self, cv_text: str, output_path: Path):
        """Generate DOCX with clean formatting."""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        lines = cv_text.split('\n')
        
        for line in lines:
            if not line.strip():
                doc.add_paragraph()  # Blank line
                continue
            
            # Headers
            if line.isupper() and len(line) < 50:
                p = doc.add_paragraph(line)
                p.style = 'Heading 1'
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
            # Bullets
            elif line.strip().startswith('•'):
                p = doc.add_paragraph(line, style='List Bullet')
                p.runs[0].font.size = Pt(11)
            # Regular text
            else:
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(11)
        
        doc.save(str(output_path))

class HTMLExporter:
    """Export CV to HTML for web preview."""
    
    def export(self, cv_text: str, output_path: Path):
        """Generate HTML with Bootstrap styling."""
        # Convert text to HTML with basic formatting
        html_body = self._text_to_html(cv_text)
        
        html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Optimized CV</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
        .cv-container {{ max-width: 800px; margin: 2rem auto; padding: 2rem; background: white; }}
        h1 {{ font-size: 1.5rem; font-weight: bold; margin-bottom: 0.5rem; }}
        h2 {{ font-size: 1.2rem; font-weight: bold; margin-top: 1.5rem; margin-bottom: 0.5rem; border-bottom: 2px solid #333; }}
        ul {{ list-style-type: disc; padding-left: 1.5rem; }}
        .contact-info {{ margin-bottom: 1rem; color: #555; }}
    </style>
</head>
<body>
    <div class="cv-container">
        {html_body}
    </div>
</body>
</html>
        """
        
        output_path.write_text(html_template, encoding='utf-8')
    
    def _text_to_html(self, text: str) -> str:
        """Convert plain text CV to HTML."""
        lines = text.split('\n')
        html_lines = []
        in_list = False
        
        for line in lines:
            if not line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append('<br>')
                continue
            
            # Headers
            if line.isupper() and len(line) < 50:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h2>{line}</h2>')
            # Bullets
            elif line.strip().startswith('•'):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                html_lines.append(f'<li>{line.strip()[1:].strip()}</li>')
            # Regular text
            else:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<p>{line}</p>')
        
        if in_list:
            html_lines.append('</ul>')
        
        return '\n'.join(html_lines)

class JSONExporter:
    """Export structured CV to JSON."""
    
    def export(self, cv_dict: Dict, output_path: Path):
        """Save CV as JSON with metadata."""
        output = {
            "exported_at": datetime.now().isoformat(),
            "format_version": "1.0",
            "cv": cv_dict
        }
        
        output_path.write_text(
            json.dumps(output, indent=2, ensure_ascii=False),
            encoding='utf-8'
        )
